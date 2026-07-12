import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from langchain.tools import tool
from pypdf import PdfReader


PROJECT_ROOT = Path(__file__).resolve().parents[1]
MATERIALS_DIR = PROJECT_ROOT / "materials"
SKILLS_DIR = PROJECT_ROOT / "skills"
OUTPUT_DIR = PROJECT_ROOT / "output"
SKILL_DETAIL_FILE = SKILLS_DIR / "skills_detail.md"

SUPPORTED_MATERIAL_SUFFIXES = {".md", ".txt", ".docx", ".pdf"}
MAX_MATERIAL_CHARS = 20000
MAX_SKILL_CHARS = 30000


def _safe_material_path(filename: str) -> Path:
    """解析用户传入的文件名，并限制只能访问 materials 文件夹。"""
    file_path = (MATERIALS_DIR / filename).resolve()
    materials_root = MATERIALS_DIR.resolve()
    if materials_root not in file_path.parents and file_path != materials_root:
        raise ValueError("只能读取 materials 文件夹内部的文件。")
    return file_path


def _safe_skill_path(filename: str) -> Path:
    """解析 skill 文件名，并限制只能读取 skills 文件夹内部的 md 文件。"""
    file_path = (SKILLS_DIR / filename).resolve()
    skills_root = SKILLS_DIR.resolve()
    if skills_root not in file_path.parents and file_path != skills_root:
        raise ValueError("只能读取 skills 文件夹内部的文件。")
    if file_path.suffix.lower() != ".md":
        raise ValueError("只能读取 Markdown 格式的 skill 文件。")
    return file_path


def _read_text_file(file_path: Path) -> str:
    """按常见中文编码读取文本文件。"""
    for encoding in ("utf-8", "utf-8-sig", "gbk"):
        try:
            return file_path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise UnicodeDecodeError("unknown", b"", 0, 1, "无法用 UTF-8 或 GBK 解码该文本文件")


def _read_docx_file(file_path: Path) -> str:
    """提取 docx 文件中的段落和表格文本。"""
    document = Document(str(file_path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    table_lines = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_lines.append(" | ".join(cells))

    return "\n".join(paragraphs + table_lines)


def _read_pdf_file(file_path: Path) -> str:
    """提取文字型 PDF 中的文本；扫描件 PDF 通常无法提取。"""
    reader = PdfReader(str(file_path))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[第{index}页]\n{text.strip()}")
    return "\n\n".join(pages)


def _clean_report_filename(filename: str) -> str:
    """清理文件名中的非法字符，并保证报告使用 md 后缀。"""
    cleaned = re.sub(r'[\\/:*?"<>|]', "_", filename).strip()
    cleaned = cleaned.rstrip(". ")
    if not cleaned:
        cleaned = f"法律报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    if not cleaned.lower().endswith(".md"):
        cleaned += ".md"
    return cleaned


@tool
def list_material_files() -> str:
    """
    列出 materials 文件夹中可读取的法律材料文件。

    返回：
        JSON 字符串，包含文件名、后缀和大小。仅列出 .md、.txt、.docx、文字型 .pdf。
    """
    print("正在查看材料文件")
    MATERIALS_DIR.mkdir(parents=True, exist_ok=True)
    files = []
    for file_path in sorted(MATERIALS_DIR.iterdir()):
        if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_MATERIAL_SUFFIXES:
            files.append({
                "filename": file_path.name,
                "suffix": file_path.suffix.lower(),
                "size_bytes": file_path.stat().st_size,
            })
    return json.dumps(files, ensure_ascii=False)


@tool
def read_legal_material(filename: str) -> str:
    """
    读取 materials 文件夹中的法律材料，并提取为纯文本。

    参数：
        filename: materials 文件夹内的文件名，例如 contract.md、borrow_note.txt、judgment.docx、case.pdf。

    返回：
        JSON 字符串，包含文件名、文件类型、提取文本和提示信息。
    """
    print("正在读取材料")
    try:
        file_path = _safe_material_path(filename)
        if not file_path.exists():
            return json.dumps({"error": f"文件不存在：{filename}"}, ensure_ascii=False)
        if not file_path.is_file():
            return json.dumps({"error": f"路径不是文件：{filename}"}, ensure_ascii=False)

        suffix = file_path.suffix.lower()
        if suffix not in SUPPORTED_MATERIAL_SUFFIXES:
            return json.dumps({
                "error": "暂不支持该文件类型。",
                "supported_suffixes": sorted(SUPPORTED_MATERIAL_SUFFIXES),
            }, ensure_ascii=False)

        if suffix in {".md", ".txt"}:
            text = _read_text_file(file_path)
        elif suffix == ".docx":
            text = _read_docx_file(file_path)
        else:
            text = _read_pdf_file(file_path)

        text = text.strip()
        if not text:
            message = "未提取到文字内容。若这是扫描版 PDF 或图片文件，当前工具不支持 OCR。"
        else:
            message = "提取成功。"

        truncated = len(text) > MAX_MATERIAL_CHARS
        if truncated:
            text = text[:MAX_MATERIAL_CHARS]

        return json.dumps({
            "filename": file_path.name,
            "suffix": suffix,
            "content": text,
            "truncated": truncated,
            "message": message,
        }, ensure_ascii=False)

    except Exception as exc:
        return json.dumps({"error": f"读取材料失败：{exc}"}, ensure_ascii=False)


@tool
def list_report_skills() -> str:
    """
    读取 skills_detail.md，查看当前支持的报告和文书生成 skill。

    返回：
        Markdown 文本，包含可用 skill、对应文件和用途说明。
    """
    print("正在查看有那些SKILL")
    if not SKILL_DETAIL_FILE.exists():
        return "未找到 skill 目录文件：skills/skills_detail.md"
    return SKILL_DETAIL_FILE.read_text(encoding="utf-8").strip()


@tool
def read_report_skill(skill_file: str) -> str:
    """
    读取某一个具体的报告或文书生成 skill 文件。

    参数：
        skill_file: skills 文件夹内的 skill 文件名，例如 loan_dispute_report.md。

    返回：
        Markdown 文本，包含该 skill 的适用场景、建议结构和写作要求。
    """
    print("正在读取SKILL")
    try:
        skill_path = _safe_skill_path(skill_file)
        if not skill_path.exists():
            return f"未找到 skill 文件：{skill_file}"
        text = skill_path.read_text(encoding="utf-8").strip()
        if len(text) > MAX_SKILL_CHARS:
            return text[:MAX_SKILL_CHARS]
        return text
    except Exception as exc:
        return f"读取 skill 失败：{exc}"


@tool
def save_markdown_report(filename: str, content: str) -> str:
    """
    将 Markdown 报告或法律文书草稿保存到 output 文件夹。

    参数：
        filename: 中文文件名，例如 借贷纠纷分析报告.md、合同审查报告.md。
        content: Markdown 正文内容。

    返回：
        JSON 字符串，包含保存状态和报告路径。
    """
    print("正在生成报告")
    try:
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        safe_name = _clean_report_filename(filename)
        report_path = OUTPUT_DIR / safe_name

        if report_path.exists():
            stem = report_path.stem
            suffix = report_path.suffix
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            report_path = OUTPUT_DIR / f"{stem}_{timestamp}{suffix}"

        report_path.write_text(content, encoding="utf-8")
        return json.dumps({
            "saved": True,
            "filename": report_path.name,
            "path": str(report_path),
        }, ensure_ascii=False)

    except Exception as exc:
        return json.dumps({"saved": False, "error": f"保存报告失败：{exc}"}, ensure_ascii=False)


@tool
def search_legal_statutes(query: str, n_results: int) -> str:
    """
    根据查询内容检索相关法律条文。

    参数：
        query: 法律关键词或问题。
        n_results: 返回结果数量，最多 10 个。

    返回：
        JSON 字符串，包含匹配的法律条文内容。
    """
    from rag.Rag import corpus_query

    n_results = min(n_results, 10)
    results = corpus_query(query=query, n_results=n_results)
    print("正在使用法条检索工具", query, n_results)
    return json.dumps(results, ensure_ascii=False)


@tool
def search_legal_QA(query: str, n_results: int) -> str:
    """
    根据查询内容检索相关法律问答案例。

    参数：
        query: 法律问题或具体场景。
        n_results: 返回结果数量，最多 10 个。

    返回：
        JSON 字符串，包含匹配的法律问答内容。
    """
    from rag.Rag import qa_query

    n_results = min(n_results, 10)
    results = qa_query(query=query, n_results=n_results)
    print("正在使用法律问答检索工具", query, n_results)
    return json.dumps(results, ensure_ascii=False)
